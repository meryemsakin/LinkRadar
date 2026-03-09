from __future__ import annotations
"""
Word Parser Tool
python-docx ile Word dosyalarından yapısal metadata çıkarır.
Başlıklar, ilk paragraflar ve tablo özeti döndürür.
"""

import io
import logging

logger = logging.getLogger(__name__)


async def parse_word(file_bytes: bytes, filename: str) -> dict:
    """
    Word dosyasını (docx) parse eder. Başlıkları (Heading 1, 2, 3),
    tablo başlıklarını ve özet için ilk paragrafları çıkarır.
    Tüm metni değil, yapısal öğeleri döndürür.

    Args:
        file_bytes: Dosyanın ham byte içeriği
        filename: Dosya adı (loglama için)

    Returns:
        dict with keys:
            success: bool
            headings: list[dict] — level ve text
            intro_paragraphs: list[str] — İlk 5 paragrafın 200 char'lık kesimleri
            tables: list[dict] — Tablo başlıkları ve satır sayıları
            total_paragraphs: int
            error: str | None
    """
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))

        headings = []
        first_paragraphs = []
        tables_summary = []

        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level_str = para.style.name.replace("Heading ", "").strip()
                level = int(level_str) if level_str.isdigit() else 1
                if para.text.strip():
                    headings.append({
                        "level": level,
                        "text": para.text.strip(),
                    })
            elif para.text.strip() and len(first_paragraphs) < 5:
                first_paragraphs.append(para.text.strip()[:200])

        for i, table in enumerate(doc.tables[:5]):  # Max 5 tablo
            if table.rows:
                header_row = [cell.text.strip() for cell in table.rows[0].cells]
                tables_summary.append({
                    "table_index": i + 1,
                    "headers": header_row,
                    "row_count": len(table.rows),
                })

        logger.info(
            f"Word parse başarılı: {filename} — "
            f"{len(headings)} başlık, {len(tables_summary)} tablo"
        )

        return {
            "success": True,
            "headings": headings,
            "intro_paragraphs": first_paragraphs,
            "tables": tables_summary,
            "total_paragraphs": len(doc.paragraphs),
            "error": None,
        }

    except Exception as e:
        logger.error(f"Word parse hatası: {filename} — {str(e)}")
        return {
            "success": False,
            "headings": [],
            "intro_paragraphs": [],
            "tables": [],
            "total_paragraphs": 0,
            "error": f"Word parse hatası: {str(e)}",
        }
